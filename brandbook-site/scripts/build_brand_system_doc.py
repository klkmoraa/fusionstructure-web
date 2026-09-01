from __future__ import annotations

from pathlib import Path
from shutil import copy2
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path(
    r"C:\Users\crisd\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-system-design\assets\reference.docx"
)
OUTPUT = ROOT / "docs" / "FusionStructure-Brand-System-Architecture.docx"


def set_table(table, rows):
    for row_index, values in enumerate(rows):
        if row_index >= len(table.rows):
            break
        for cell_index, value in enumerate(values):
            if cell_index < len(table.rows[row_index].cells):
                table.rows[row_index].cells[cell_index].text = value


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    copy2(REFERENCE, OUTPUT)
    document = Document(OUTPUT)
    paragraphs = document.paragraphs

    replacements = {
        8: "FusionStructure",
        9: "Arquitectura del sistema visual",
        21: "1.  Resumen",
        22: (
            "FusionStructure adopta un sistema visual propio para convertir modelos, resultados y decisiones "
            "estructurales en una experiencia clara, conectada y trazable. Su principio rector es Make complexity "
            "legible.: la geometría comunica estructura, el color comunica significado y el movimiento explica "
            "causa y efecto."
        ),
        23: (
            "El sistema cubre identidad, tokens, superficies, controles, iconos, diagramas, copy, movimiento y "
            "comportamiento responsive. Describe una dirección de producto; no certifica cálculos, normativa ni "
            "preparación para obra."
        ),
        25: "2.  Objetivos y no objetivos",
        27: "3.  Contexto y problema",
        28: (
            "La interfaz debe mostrar información técnica densa sin ocultar texto, perder jerarquía o depender de "
            "un único color. La versión anterior permitía que el canvas invadiera el titular y usaba una superficie "
            "nocturna demasiado negra, con una sombra casi imperceptible."
        ),
        29: (
            "El límite del sistema visual comienza en tokens semánticos y termina en componentes renderizados. "
            "Los datos estructurales siguen perteneciendo al producto y sus solvers; el brandbook sólo define cómo "
            "representarlos con consistencia."
        ),
        30: "4.  Arquitectura propuesta",
        33: "Figura 1. Capas del sistema: intención → señal → componente → superficie → viewport.",
        35: "Componentes centrales",
        39: "5.  Ciclo de una interacción",
        40: "1. La vista recibe contenido, estado, tema, viewport y preferencia de movimiento.",
        41: "2. Se valida que la señal semántica corresponda al resultado o estado mostrado.",
        42: "3. Se cargan tokens de superficie, tipografía, espaciado, borde, sombra y movimiento.",
        43: "4. El componente selecciona jerarquía, icono, copy y diagrama sin alterar el dato fuente.",
        44: "5. El estado persistente se guarda en el dominio del producto, no en el componente visual.",
        45: "6. La transición usa 180–520 ms, spring suave y una alternativa sin movimiento.",
        46: "7. La salida se valida por legibilidad, contraste, foco, overflow y coherencia entre temas.",
        48: "6.  Contratos de diseño",
        49: "Contrato primario de tokens",
        52: "Garantías del contrato",
        54: "Los nombres semánticos son estables aunque cambie el valor visual entre día y noche.",
        55: "Cada variante y estado tiene un identificador único y una prioridad explícita.",
        56: "Tema, viewport, estado, preferencia de movimiento y versión del sistema quedan trazables.",
        57: "Los tokens son fuente de verdad visual; no son fuente de verdad para resultados estructurales.",
        58: "La implementación versionada vive en el Site del brandbook y en sus assets de marca.",
        59: "Los cambios deben actualizar el Site, este documento y las pruebas visuales relacionadas.",
        61: "7. Consistencia, repetición y recuperación",
        63: (
            "La misma combinación de señal, estado y tema debe producir la misma jerarquía. Repetir una interacción "
            "no debe duplicar efectos. Si faltan tokens o assets, el componente usa una superficie neutral, conserva "
            "el contenido y expone el estado degradado."
        ),
        65: "8. Accesibilidad, seguridad y privacidad",
        67: "Los componentes respetan foco visible, orden de teclado y objetivos táctiles de al menos 44 px.",
        68: "El color nunca es el único portador de significado; icono, etiqueta y posición lo acompañan.",
        69: "El brandbook no almacena credenciales ni datos de proyecto; los mockups usan contenido ficticio.",
        70: "El modo de movimiento reducido elimina desplazamientos y conserva cambios de estado instantáneos.",
        71: "El copy evita promesas de certificación, exactitud normativa o preparación para obra sin evidencia.",
        74: "9.  Preparación operativa",
        76: "10. Alternativas consideradas",
        80: "11. Preguntas abiertas",
        81: "1. ¿Qué secuencia de adopción seguirá la nueva marca madre y su familia de herramientas?",
        82: "2. ¿Qué diagramas serán Disponible, Experimental y Planeado en la primera adopción?",
        83: "3. ¿El tema seguirá al sistema operativo o será una preferencia guardada por proyecto?",
        84: "4. ¿Qué equipo aprueba cambios futuros en tokens semánticos y copy crítico?",
        86: "12. Decisión y siguientes pasos",
        87: (
            "Adoptar este sistema como dirección visual propuesta. Primero se valida el Site responsive y su "
            "paridad día/noche; después se migran componentes de producto por familias reversibles, empezando por "
            "navegación, tarjetas y diagramas. La adopción amplia requiere contraste, teclado, movimiento reducido, "
            "pruebas visuales y límites de producto explícitos."
        ),
    }
    for index, text in replacements.items():
        paragraphs[index].text = text

    set_table(document.tables[0], [["ESTADO\nPropuesto", "", "PROPIETARIO\nFusionStructure", "", "ACTUALIZADO\nSeptiembre 1, 2026"]])
    set_table(
        document.tables[1],
        [
            ["Autores", "Codex + FusionStructure"],
            ["Revisores", "Producto, diseño y desarrollo"],
            ["Documentos relacionados", "Brandbook interactivo de FusionStructure"],
            ["Alcance", "Identidad, componentes, movimiento, copy, temas y responsive."],
        ],
    )
    set_table(
        document.tables[2],
        [
            ["Objetivos", "No objetivos"],
            ["Titular y contenido siempre legibles.", "Rediseñar el solver o sus contratos."],
            ["Semántica estable en día y noche.", "Usar color como único significado."],
            ["Componentes interactivos y accesibles.", "Imitar la identidad de otro producto."],
            ["Migración gradual y reversible.", "Afirmar certificación o uso profesional."],
        ],
    )
    set_table(
        document.tables[3],
        [
            ["Componente", "Responsabilidad", "Fuente", "Fallo seguro"],
            ["Tokens", "Color, tipo, espacio, borde y sombra.", "CSS variables", "Neutraliza sin ocultar contenido."],
            ["Marca", "Identidad y reconocimiento.", "SVG versionado", "Wordmark textual como fallback."],
            ["Señales", "Mapear significado técnico a color.", "Mapa semántico", "Etiqueta e icono siguen visibles."],
            ["Componentes", "Componer controles, tarjetas y estados.", "React/CSS", "Estado explícito y reversible."],
            ["QA", "Validar viewport, tema, foco y motion.", "Checks + navegador", "Bloquea publicación defectuosa."],
        ],
    )
    set_table(
        document.tables[4],
        [
            ["Campo", "Tipo", "Requerido", "Descripción"],
            ["intent", "string", "Sí", "Acción o información que el componente comunica."],
            ["signal", "enum", "Sí", "Axial, momento, cortante, deformada, fluencia o estado."],
            ["theme", "enum", "Sí", "Día o noche con igual semántica."],
            ["state", "enum", "Sí", "Default, hover, focus, active, loading, disabled o error."],
            ["viewport", "number", "Sí", "Ancho disponible; nunca ancho de pantalla supuesto."],
            ["motion", "enum", "Sí", "Completo o reducido."],
            ["version", "string", "Sí", "Versión trazable del sistema visual."],
        ],
    )
    set_table(
        document.tables[5],
        [
            ["Escenario", "Comportamiento esperado", "Razón"],
            ["Interacción repetida", "Mismo estado, sin efectos duplicados.", "Consistencia perceptual."],
            ["Token o asset ausente", "Fallback neutral y contenido intacto.", "La legibilidad tiene prioridad."],
            ["Viewport estrecho", "Reflujo a una columna, sin overflow.", "Paridad móvil."],
            ["Cambio de tema", "Geometría estable; cambian papel y profundidad.", "Reduce desorientación."],
        ],
    )
    set_table(
        document.tables[6],
        [
            ["Señal", "Objetivo", "Responsable", "Puerta"],
            ["Legibilidad", "Cero texto oculto o solapado.", "Diseño", "Requerida"],
            ["Responsive", "Cero overflow a 390 px.", "Frontend", "Requerida"],
            ["Contraste", "Estados y foco discernibles.", "Diseño + QA", "Requerida"],
            ["Movimiento", "Sin mareo; reduced motion funcional.", "Frontend", "Requerida"],
            ["Build", "Lint, build y checks verdes.", "Desarrollo", "Requerida"],
            ["Restricción de rollout: publicar primero el brandbook; migrar el producto por familias reversibles."] * 4,
        ],
    )
    set_table(
        document.tables[7],
        [
            ["Alternativa", "Motivo", "Por qué no se elige"],
            ["Minimalismo plano", "Alta simplicidad inicial.", "Pierde profundidad e interacción."],
            ["Oscuro negro puro", "Contraste aparente.", "Aplana superficies y esconde sombras."],
            ["Color decorativo", "Impacto visual rápido.", "Debilita el significado técnico."],
            ["Copiar otra marca", "Reduce tiempo de exploración.", "No crea identidad propia."],
        ],
    )
    set_table(
        document.tables[8],
        [
            ["Hito", "Entregable", "Criterio de salida"],
            ["M1", "Brandbook responsive y logo.", "Desktop/móvil, día/noche y copy validados."],
            ["M2", "Tokens y componentes compartidos.", "Pruebas de estados y accesibilidad."],
            ["M3", "Primera familia dentro del producto.", "Adopción reversible sin regresiones."],
            ["M4", "Expansión gradual.", "Paridad y trazabilidad verificadas."],
        ],
    )

    document.core_properties.title = "FusionStructure — Arquitectura del sistema visual"
    document.core_properties.subject = "Make complexity legible."
    document.core_properties.author = "FusionStructure + Codex"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
